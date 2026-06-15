"""
מערכת צ'אט משפטי לאסתי - בית דין רבני
=============================================
שימוש:
    pip install anthropic python-docx

    set ANTHROPIC_API_KEY=sk-ant-...  (Windows)
    export ANTHROPIC_API_KEY=sk-ant-...  (Linux/Mac)

    python legal_chat.py
    python legal_chat.py --docs D:\\12  (ספרייה עם המסמכים)
"""

import os
import sys
import json
import argparse
import datetime
from pathlib import Path

try:
    import anthropic
except ImportError:
    print("ERROR: pip install anthropic")
    sys.exit(1)

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("אזהרה: python-docx לא מותקן. ייצוא Word לא יהיה זמין.")
    print("הרץ: pip install python-docx")

# ─── הגדרות ─────────────────────────────────────────────────────────────────

DOCS_ROOT = Path(r"D:\12")
OUTPUT_DIR = Path(r"D:\12\outputs")

SYSTEM_PROMPT = """אתה עוזר משפטי מומחה לדיני משפחה ובית דין רבני בישראל.
אתה עוזר לאסתי הלר-גפנר בניהול התיק שלה מול בעלה לשעבר יוסף גפנר.

הרקע המרכזי:
- הצדדים נשואים ובהליכי גירושין בבית הדין הרבני
- יש פסיקתא מיום 29/03/2026 המורה על פינוי יוסף מהדירה
- הדירה רשומה בטאבו על שם אסתי
- יוסף לא הגיש מסמכי חובות מאז יוני 2025 (כשנה)
- ההחלטה מ-03/06/2026 עוסקת בחובות בלבד ואינה מבטלת את צו הפינוי
- מתועדת אלימות ואיומים על נציגים משפטיים
- יש 4 ילדים משותפים

כישורים שלך:
1. ניתוח מסמכים משפטיים בעברית
2. ניסוח בקשות, עתירות ומכתבים לבית דין רבני
3. הסבר הליכים ומשמעות החלטות
4. יצירת טפסים ובקשות מובנות
5. ייצוא מסמכים לפורמט Word

כללי עבודה:
- תמיד דבר בעברית ברורה ומובנת
- הסבר מונחים משפטיים בשפה פשוטה
- ציין מתי נדרשת התייעצות עם עורך דין
- שמור על סודיות ורגישות הנושאים
- כשמייצר מסמך, הוסף תאריך ומספר תיק

כשאתה מייצר מסמך רשמי:
- השתמש בכלי generate_document
- כלול כותרת, תאריך, שם המגישה, וגוף המסמך
- ניסוח בשפה משפטית מתאימה לבית דין רבני"""

# ─── כלים ────────────────────────────────────────────────────────────────────

TOOLS = [
    {
        "name": "generate_document",
        "description": "מייצר מסמך Word עם תוכן משפטי ושומר אותו לתיקיית הפלט. השתמש בכלי זה כשהמשתמש מבקש ליצור בקשה, עתירה, מכתב או כל מסמך רשמי.",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "כותרת המסמך (למשל: 'בקשה לאכיפת צו פינוי')"
                },
                "document_type": {
                    "type": "string",
                    "enum": ["בקשה", "עתירה", "מכתב", "תצהיר", "הודעה", "ערעור"],
                    "description": "סוג המסמך"
                },
                "to": {
                    "type": "string",
                    "description": "למי המסמך מיועד (למשל: 'כבוד בית הדין הרבני האזורי')"
                },
                "subject": {
                    "type": "string",
                    "description": "נושא המסמך"
                },
                "body": {
                    "type": "string",
                    "description": "גוף המסמך המלא - הפסקאות, הנימוקים, הבקשות"
                },
                "requests": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "רשימת הבקשות/הסעדים המבוקשים (נקודות)"
                },
                "filename": {
                    "type": "string",
                    "description": "שם הקובץ לשמירה (ללא סיומת, עברית/אנגלית)"
                }
            },
            "required": ["title", "document_type", "body", "filename"]
        }
    },
    {
        "name": "search_case_files",
        "description": "מחפש בקבצי התיק המקומיים לפי מילות מפתח. מחזיר רשימת קבצים רלוונטיים.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "מילות חיפוש (עברית או אנגלית)"
                },
                "file_types": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "סוגי קבצים לחיפוש (למשל: ['.pdf', '.docx']). ריק = כל הסוגים"
                }
            },
            "required": ["query"]
        }
    },
    {
        "name": "read_case_file",
        "description": "קורא תוכן של קובץ טקסט מתיק התיק",
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "נתיב מלא לקובץ"
                }
            },
            "required": ["file_path"]
        }
    },
    {
        "name": "list_case_files",
        "description": "מציג רשימת קבצים בתיק (או בתיקיית משנה)",
        "input_schema": {
            "type": "object",
            "properties": {
                "subfolder": {
                    "type": "string",
                    "description": "תיקיית משנה לחיפוש (אופציונלי - ריק לכל התיק)"
                }
            },
            "required": []
        }
    }
]

# ─── ביצוע כלים ──────────────────────────────────────────────────────────────

def run_generate_document(args: dict) -> str:
    if not DOCX_AVAILABLE:
        return "שגיאה: python-docx לא מותקן. הרץ: pip install python-docx"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()

    # כותרת ראשית
    doc.add_heading(args.get("document_type", "מסמך"), level=1)
    doc.add_heading(args["title"], level=2)

    # פרטי כותרת
    today = datetime.date.today().strftime("%d/%m/%Y")
    doc.add_paragraph(f"תאריך: {today}")
    if args.get("to"):
        doc.add_paragraph(f"לכבוד: {args['to']}")
    doc.add_paragraph("הנדון: " + args.get("subject", args["title"]))
    doc.add_paragraph("")

    # שם המגישה
    doc.add_paragraph("המבקשת: אסתי הלר-גפנר")
    doc.add_paragraph("המשיב: יוסף גפנר")
    doc.add_paragraph("")

    # גוף המסמך
    for para in args["body"].split("\n"):
        if para.strip():
            p = doc.add_paragraph(para.strip())

    # בקשות
    if args.get("requests"):
        doc.add_paragraph("")
        doc.add_heading("לפיכך מתבקש בית הדין הנכבד:", level=3)
        for i, req in enumerate(args["requests"], 1):
            doc.add_paragraph(f"{i}. {req}", style="List Number")

    # חתימה
    doc.add_paragraph("")
    doc.add_paragraph("בכבוד רב,")
    doc.add_paragraph("אסתי הלר-גפנר")
    doc.add_paragraph(f"תאריך: {today}")

    filename = args["filename"].replace("/", "_").replace("\\", "_")
    if not filename.endswith(".docx"):
        filename += ".docx"
    out_path = OUTPUT_DIR / filename

    doc.save(str(out_path))
    return f"✅ המסמך נשמר: {out_path}\n\nניתן לפתוח אותו ב-Microsoft Word."


def run_search_case_files(args: dict) -> str:
    query = args["query"].lower()
    file_types = args.get("file_types", [])
    root = DOCS_ROOT

    if not root.exists():
        return f"תיקיית התיק לא נמצאה: {root}\nהפעל עם --docs <נתיב>"

    results = []
    for f in root.rglob("*"):
        if not f.is_file():
            continue
        if file_types and f.suffix.lower() not in file_types:
            continue
        name_lower = f.name.lower()
        if any(w in name_lower for w in query.split()):
            rel = f.relative_to(root)
            results.append(str(rel))

    if not results:
        return f"לא נמצאו קבצים עבור: '{query}'"

    lines = [f"נמצאו {len(results)} קבצים:"]
    for r in results[:30]:
        lines.append(f"  • {r}")
    if len(results) > 30:
        lines.append(f"  ... ועוד {len(results)-30} קבצים")
    return "\n".join(lines)


def run_read_case_file(args: dict) -> str:
    path = Path(args["file_path"])
    if not path.exists():
        alt = DOCS_ROOT / args["file_path"]
        if alt.exists():
            path = alt
        else:
            return f"קובץ לא נמצא: {args['file_path']}"

    suffix = path.suffix.lower()
    if suffix == ".txt":
        try:
            return path.read_text(encoding="utf-8", errors="replace")[:8000]
        except Exception as e:
            return f"שגיאה בקריאה: {e}"
    elif suffix == ".docx" and DOCX_AVAILABLE:
        try:
            doc = DocxDocument(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text[:8000]
        except Exception as e:
            return f"שגיאה בקריאת docx: {e}"
    else:
        return f"קובץ מסוג {suffix} — ניתן לקרוא ישירות רק .txt ו-.docx. נתיב: {path}"


def run_list_case_files(args: dict) -> str:
    sub = args.get("subfolder", "")
    root = DOCS_ROOT / sub if sub else DOCS_ROOT

    if not root.exists():
        return f"תיקייה לא נמצאה: {root}"

    entries = sorted(root.iterdir(), key=lambda x: (x.is_file(), x.name))
    lines = [f"תוכן {root}:", ""]
    folders = [e for e in entries if e.is_dir()]
    files = [e for e in entries if e.is_file()]

    if folders:
        lines.append("📁 תיקיות:")
        for d in folders:
            count = sum(1 for _ in d.rglob("*") if _.is_file())
            lines.append(f"  📁 {d.name}/ ({count} קבצים)")

    if files:
        lines.append("")
        lines.append("📄 קבצים:")
        for f in files[:50]:
            size = f.stat().st_size
            sz_str = f"{size//1024}KB" if size > 1024 else f"{size}B"
            lines.append(f"  📄 {f.name} ({sz_str})")
        if len(files) > 50:
            lines.append(f"  ... ועוד {len(files)-50} קבצים")

    return "\n".join(lines)


def execute_tool(tool_name: str, tool_input: dict) -> str:
    if tool_name == "generate_document":
        return run_generate_document(tool_input)
    elif tool_name == "search_case_files":
        return run_search_case_files(tool_input)
    elif tool_name == "read_case_file":
        return run_read_case_file(tool_input)
    elif tool_name == "list_case_files":
        return run_list_case_files(tool_input)
    else:
        return f"כלי לא מוכר: {tool_name}"


# ─── לולאת שיחה ──────────────────────────────────────────────────────────────

def chat_loop(client: anthropic.Anthropic):
    messages = []

    print()
    print("=" * 60)
    print("  מערכת ייעוץ משפטי - תיק הלר גפנר")
    print("  בית הדין הרבני")
    print("=" * 60)
    print()
    print("שלום אסתי! אני כאן לעזור לך בניהול התיק.")
    print("תוכלי לשאול אותי על:")
    print("  • מה קורה בתיק שלך ומה הצעדים הבאים")
    print("  • הכנת בקשות ומסמכים לבית הדין")
    print("  • הסבר החלטות ומושגים משפטיים")
    print("  • חיפוש בקבצי התיק")
    print()
    print("הקלידי 'יציאה' לסיום | 'נקה' לפתיחת שיחה חדשה")
    print("-" * 60)
    print()

    while True:
        try:
            user_input = input("אסתי: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n\nשלום! בהצלחה בתיק.")
            break

        if not user_input:
            continue
        if user_input in ("יציאה", "exit", "quit", "q"):
            print("שלום! בהצלחה!")
            break
        if user_input in ("נקה", "clear", "חדש"):
            messages = []
            print("\n--- שיחה חדשה ---\n")
            continue

        messages.append({"role": "user", "content": user_input})

        while True:
            response = client.messages.create(
                model="claude-opus-4-8",
                max_tokens=8096,
                thinking={"type": "adaptive"},
                system=SYSTEM_PROMPT,
                tools=TOOLS,
                messages=messages,
            )

            # אסוף תוכן
            assistant_content = response.content
            messages.append({"role": "assistant", "content": assistant_content})

            # הדפס טקסט
            text_parts = []
            tool_uses = []
            for block in assistant_content:
                if block.type == "text":
                    text_parts.append(block.text)
                elif block.type == "tool_use":
                    tool_uses.append(block)

            if text_parts:
                print()
                print("עוזר: " + "\n".join(text_parts))

            if response.stop_reason == "end_turn" or not tool_uses:
                print()
                break

            # בצע כלים
            tool_results = []
            for tu in tool_uses:
                print(f"\n⚙️  מפעיל: {tu.name}...")
                result = execute_tool(tu.name, tu.input)
                print(result[:300] + ("..." if len(result) > 300 else ""))
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tu.id,
                    "content": result
                })

            messages.append({"role": "user", "content": tool_results})


# ─── נקודת כניסה ─────────────────────────────────────────────────────────────

def main():
    global DOCS_ROOT, OUTPUT_DIR

    parser = argparse.ArgumentParser(description="מערכת צ'אט משפטי לתיק הלר-גפנר")
    parser.add_argument("--docs", type=str, default=str(DOCS_ROOT),
                        help=f"נתיב לתיקיית התיק (ברירת מחדל: {DOCS_ROOT})")
    parser.add_argument("--output", type=str, default=str(OUTPUT_DIR),
                        help="תיקיית פלט למסמכים")
    args = parser.parse_args()

    DOCS_ROOT = Path(args.docs)
    OUTPUT_DIR = Path(args.output)

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: יש להגדיר משתנה סביבה ANTHROPIC_API_KEY")
        print()
        print("Windows:")
        print('  set ANTHROPIC_API_KEY=sk-ant-...')
        print()
        print("Mac/Linux:")
        print('  export ANTHROPIC_API_KEY=sk-ant-...')
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)

    if not DOCS_ROOT.exists():
        print(f"אזהרה: תיקיית התיק לא נמצאה: {DOCS_ROOT}")
        print("חיפוש בקבצים לא יעבוד, אך שאר הפונקציות זמינות.")
        print()

    chat_loop(client)


if __name__ == "__main__":
    main()
